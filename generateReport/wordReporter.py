from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_word_report(data, output_path, title="数据分析报告"):
    """生成Word格式报告"""
    doc = Document()

    # 添加标题
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加表格
    table = doc.add_table(rows=1, cols=len(data.columns))
    table.style = 'Light Shading Accent 1'

    # 添加表头
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(data.columns):
        hdr_cells[i].text = str(column)

    # 添加数据行
    for _, row in data.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    # 添加图表说明
    doc.add_paragraph("\n数据统计结果:", style='Heading 2')

    # 保存文档
    doc.save(output_path)
    print(f"Word报告已生成: {output_path}")